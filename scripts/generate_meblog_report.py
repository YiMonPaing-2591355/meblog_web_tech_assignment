from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report_output"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "7CC005_2532764_Hein_Zarni_Tun_Web_Technologies_Report_MeBlog.docx"

PRIMARY = RGBColor(26, 68, 73)
ACCENT = RGBColor(44, 120, 107)
LIGHT = RGBColor(230, 241, 239)
TEXT = RGBColor(34, 34, 34)
MUTED = RGBColor(96, 96, 96)


def ensure_dirs():
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)


def try_font(name: str, size: int):
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts") / "arial.ttf",
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


FONT_REG = try_font("arial.ttf", 28)
FONT_BOLD = try_font("arialbd.ttf", 30)
FONT_SMALL = try_font("arial.ttf", 22)


def rounded_box(draw, xy, fill, outline, radius=18, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw, box, text, font, fill):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    x = left + (right - left - text_w) / 2
    y = top + (bottom - top - text_h) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=6, align="center")


def arrow(draw, start, end, fill=(44, 120, 107), width=5):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        direction = 1 if ex > sx else -1
        draw.polygon([(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)], fill=fill)
    else:
        direction = 1 if ey > sy else -1
        draw.polygon([(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)], fill=fill)


def make_architecture_diagram(path: Path):
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 110), fill=(26, 68, 73))
    draw.text((60, 30), "MeBlog System Architecture", font=FONT_BOLD, fill="white")
    draw.text((60, 70), "React SPA + Laravel API + MySQL + Local Media Storage", font=FONT_SMALL, fill="white")

    boxes = {
        "browser": (90, 220, 420, 430),
        "frontend": (520, 180, 900, 470),
        "backend": (1020, 180, 1430, 470),
        "database": (1040, 570, 1260, 770),
        "storage": (1300, 570, 1520, 770),
    }

    rounded_box(draw, boxes["browser"], (244, 249, 248), (44, 120, 107))
    centered_text(draw, boxes["browser"], "Browser User\n\nReaders, Authors,\nAdministrators", FONT_REG, (34, 34, 34))

    rounded_box(draw, boxes["frontend"], (230, 241, 239), (44, 120, 107))
    centered_text(
        draw,
        boxes["frontend"],
        "Frontend (React)\n\nReact Router\nAxios client\nRole-based pages\nSearch, comments,\npost creation and review",
        FONT_REG,
        (34, 34, 34),
    )

    rounded_box(draw, boxes["backend"], (232, 244, 250), (34, 87, 122))
    centered_text(
        draw,
        boxes["backend"],
        "Backend (Laravel 12 API)\n\nSanctum auth\nControllers + Requests\nRole middleware\nPost workflow and moderation",
        FONT_REG,
        (34, 34, 34),
    )

    rounded_box(draw, boxes["database"], (251, 247, 237), (180, 124, 33))
    centered_text(draw, boxes["database"], "MySQL Database\n\nusers\ncategories\nposts\ncomments", FONT_REG, (34, 34, 34))

    rounded_box(draw, boxes["storage"], (248, 240, 245), (145, 65, 103))
    centered_text(draw, boxes["storage"], "Public Media Storage\n\nuploaded cover images\nserved by /media/{path}", FONT_REG, (34, 34, 34))

    arrow(draw, (420, 325), (520, 325))
    arrow(draw, (900, 260), (1020, 260))
    arrow(draw, (1020, 385), (900, 385))
    arrow(draw, (1130, 470), (1130, 570))
    arrow(draw, (1360, 470), (1360, 570))

    draw.text((560, 520), "Communication flow: Browser -> React -> Laravel API -> Database / Media", font=FONT_SMALL, fill=(96, 96, 96))
    img.save(path)


def entity_box(draw, xy, title, fields, outline):
    rounded_box(draw, xy, (252, 252, 252), outline)
    left, top, right, bottom = xy
    draw.rectangle((left, top, right, top + 52), fill=outline)
    draw.text((left + 18, top + 12), title, font=FONT_BOLD, fill="white")
    text = "\n".join(fields)
    draw.multiline_text((left + 18, top + 72), text, font=FONT_SMALL, fill=(34, 34, 34), spacing=8)


def make_erd_diagram(path: Path):
    img = Image.new("RGB", (1600, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 110), fill=(26, 68, 73))
    draw.text((60, 30), "MeBlog Entity Relationship Overview", font=FONT_BOLD, fill="white")
    draw.text((60, 70), "Core entities used by the implemented blogging workflow", font=FONT_SMALL, fill="white")

    users = (80, 180, 450, 460)
    categories = (590, 180, 960, 420)
    posts = (1080, 150, 1500, 570)
    comments = (590, 580, 1010, 860)

    entity_box(draw, users, "users", ["PK id", "name", "email", "password", "role", "created_at"], (44, 120, 107))
    entity_box(draw, categories, "categories", ["PK id", "name", "slug", "created_at"], (180, 124, 33))
    entity_box(
        draw,
        posts,
        "posts",
        [
            "PK id",
            "FK user_id -> users.id",
            "FK category_id -> categories.id",
            "title, slug, content",
            "image, status",
            "submitted_at, approved_at",
            "FK approved_by -> users.id",
            "published_at",
        ],
        (34, 87, 122),
    )
    entity_box(
        draw,
        comments,
        "comments",
        [
            "PK id",
            "FK user_id -> users.id",
            "FK post_id -> posts.id",
            "comment",
            "status",
            "created_at",
        ],
        (145, 65, 103),
    )

    arrow(draw, (450, 290), (1080, 290))
    arrow(draw, (960, 300), (1080, 300))
    arrow(draw, (280, 460), (700, 580))
    arrow(draw, (1180, 570), (920, 720))
    draw.text((700, 455), "A user writes many posts", font=FONT_SMALL, fill=(96, 96, 96))
    draw.text((1040, 620), "A post has many comments", font=FONT_SMALL, fill=(96, 96, 96))
    draw.text((620, 455), "A category contains many posts", font=FONT_SMALL, fill=(96, 96, 96))
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=TEXT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_header_footer(section, header_text):
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.text = header_text
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p.runs:
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT

    for style_name, size in [("Title", 24), ("Subtitle", 13)]:
        style = doc.styles[style_name]
        style.font.name = "Georgia"
        style.font.color.rgb = PRIMARY
        style.font.size = Pt(size)

    if "Report Heading 1" not in doc.styles:
        style = doc.styles.add_style("Report Heading 1", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 1"]
    style = doc.styles["Report Heading 1"]
    style.font.name = "Georgia"
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = PRIMARY

    if "Report Heading 2" not in doc.styles:
        style = doc.styles.add_style("Report Heading 2", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 2"]
    style = doc.styles["Report Heading 2"]
    style.font.name = "Georgia"
    style.font.size = Pt(13)
    style.font.bold = True
    style.font.color.rgb = ACCENT

    if "Caption Custom" not in doc.styles:
        style = doc.styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles["Caption Custom"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.font.italic = True
    style.font.color.rgb = MUTED


def add_title_page(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.different_first_page_header_footer = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HEIN ZARNI TUN 7CC005_WEB_TECHNOLOGIES").bold = True
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = PRIMARY

    doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("MeBlog - Role-Based Blogging Platform")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Web Technologies Assignment Report")

    doc.add_paragraph()
    metadata = [
        ("Name", "Hein Zarni Tun"),
        ("Student ID", "2532764"),
        ("Course Title", "MSc Computer Science"),
        ("Module Name", "Web Technologies"),
        ("Module Code", "7CC005"),
        ("Center Name", "Strategy First University, Yangon, Myanmar"),
        ("Submission Date", "April 2026"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        a = p.add_run(f"{label}: ")
        a.bold = True
        a.font.color.rgb = PRIMARY
        b = p.add_run(value)
        b.font.color.rgb = TEXT

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run(
        "This report documents the design, implementation, testing, and technical discussion "
        "for MeBlog, a full-stack blogging application built with React and Laravel."
    )
    run.font.size = Pt(11)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    style = "Report Heading 1" if level == 1 else "Report Heading 2"
    p = doc.add_paragraph(style=style)
    p.space_before = Pt(8)
    p.space_after = Pt(6)
    p.add_run(text)
    return p


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        first.bold = True
        rest = p.add_run(text[len(bold_prefix):])
        rest.font.color.rgb = TEXT
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_contents(doc):
    add_heading(doc, "Table of Contents", 1)
    items = [
        "Task 1 - Implementation",
        "1.1 Introduction",
        "1.2 Tools and Technologies Used",
        "1.3 Development Process",
        "1.4 Version Control and Code Management",
        "1.5 Demonstration and Testing",
        "1.6 Challenges and Solutions",
        "1.7 Conclusion",
        "Task 2 - Testing and Discussion",
        "2.1 Testing",
        "2.2 Discussion: Comparing Single Page Applications with Traditional Web Development Methodologies",
        "References",
    ]
    add_bullets(doc, items)
    doc.add_page_break()


def insert_picture_with_caption(doc, image_path: Path, caption: str, width: float = 6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(style="Caption Custom")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)


def build_document():
    ensure_dirs()
    architecture = ASSET_DIR / "meblog_architecture.png"
    erd = ASSET_DIR / "meblog_erd.png"
    make_architecture_diagram(architecture)
    make_erd_diagram(erd)

    doc = Document()
    setup_styles(doc)
    add_title_page(doc)

    # Main section styling after cover.
    main_section = doc.sections[-1]
    add_header_footer(main_section, "HEIN ZARNI TUN 7CC005_WEB_TECHNOLOGIES")

    add_contents(doc)

    add_heading(doc, "Task 1 - Implementation", 1)
    add_heading(doc, "1.1 Introduction", 2)
    add_paragraph(
        doc,
        "Purpose. MeBlog was developed as a role-based publishing platform that allows ordinary readers to discover content, approved authors to create and submit articles, and administrators to moderate the full publishing workflow. The project demonstrates the integration of a React single-page application with a Laravel REST API, relational data modeling, authentication, moderation, and media handling."
    )
    add_paragraph(
        doc,
        "Functional Scope of the web application. The implemented system supports three primary user roles with different permissions and interfaces."
    )
    add_bullets(
        doc,
        [
            "Reader / general user: register, log in, browse published posts, search posts by keyword, filter by category, open a detailed post page, and add comments when authenticated.",
            "Author: register as an author candidate, wait for administrative approval, access an author dashboard, create posts, upload cover images, edit drafts, submit posts for review, and monitor personal post status.",
            "Administrator: approve pending authors, review pending posts, approve or reject submitted content, manage categories, moderate comments, and inspect all registered users and posts.",
        ],
    )
    add_paragraph(
        doc,
        "System Design and Modeling Diagrams. The application is organized around a decoupled frontend-backend architecture. React handles the client-side presentation and state transitions, Laravel exposes REST endpoints and business rules, and MySQL stores persistent records for users, categories, posts, and comments."
    )
    insert_picture_with_caption(doc, architecture, "Figure 1. MeBlog architecture overview")
    insert_picture_with_caption(doc, erd, "Figure 2. Core entity relationship overview")
    add_paragraph(
        doc,
        "Overview of the web application and its goals. MeBlog aims to provide a simple but structured editorial workflow instead of allowing any new author to publish instantly. This makes it suitable as an academic project because it combines public content delivery with role-based moderation, comment interaction, image upload, and administrative review in a single cohesive system."
    )

    add_heading(doc, "1.2 Tools and Technologies Used", 2)
    add_bullets(
        doc,
        [
            "Frontend: React 18, React Router DOM, and Axios were used to build the SPA, client routing, and HTTP communication.",
            "Backend: Laravel 12 and PHP 8.2 were used to implement the REST API, request validation, middleware, and moderation workflow.",
            "Authentication: Laravel Sanctum manages token-based authentication for login-protected actions.",
            "Database: MySQL stores users, categories, posts, comments, and publication state.",
            "Testing and delivery: PHPUnit verifies backend features while the Create React App build pipeline verifies the frontend bundle.",
        ],
    )
    add_paragraph(
        doc,
        "The frontend uses Create React App as the project scaffold and consumes the backend entirely through JSON APIs. Laravel Sanctum was selected because it fits the assignment well for token-based login flows without requiring an external identity provider."
    )

    add_heading(doc, "1.3 Development Process", 2)
    add_paragraph(
        doc,
        "Initial Setup. The repository was structured into two separate application folders: Frontend for the React SPA and Backend for the Laravel API. This separation made it easier to run, test, and reason about the client and server independently."
    )
    add_paragraph(
        doc,
        "Front-End Development. The user interface was organized with route-based pages for public browsing, author features, and administrative tools. Shared components such as the navigation bar, sidebar, search bar, post cards, and protected route wrappers helped keep the interface consistent across the application."
    )
    add_paragraph(
        doc,
        "Back-End Development. On the server side, Laravel controllers and request classes were used to enforce validation and permissions. The post lifecycle was implemented around draft, pending, published, and rejected statuses, while comments and categories were modeled as related entities managed through REST endpoints."
    )
    add_paragraph(
        doc,
        "Integration. Axios connects the React frontend to Laravel endpoints for registration, login, post retrieval, commenting, content submission, and moderation. During the final debugging stage, the media pipeline was refined so uploaded cover images are stored on the backend and served reliably through a dedicated /media path."
    )

    add_heading(doc, "1.4 Version Control and Code Management", 2)
    add_paragraph(
        doc,
        "Version Control and Repository Structure. Git was used throughout development, and the current repository contains separate frontend and backend folders alongside assignment-level documentation files. This arrangement reflects the overall separation of concerns of the project itself."
    )
    add_paragraph(
        doc,
        "Repository activity shows eight commits from the initial setup on 22 March 2026 through feature expansion, testing, and UI adjustments in April 2026. While the sample report demonstrates a larger number of weekly commits, this repository still records a visible progression from project creation to testing and refinement."
    )
    add_bullets(
        doc,
        [
            "22 March 2026: initial repository setup, backend and frontend created.",
            "22 March 2026: homepage and public styling adjustments applied.",
            "12 April 2026: Laravel backend structure reorganized into the main Backend directory.",
            "25 April 2026: feature test suite added and placeholder tests removed.",
            "26 April 2026: admin, authentication, homepage, registration, and image-handling improvements committed.",
        ],
    )

    add_heading(doc, "1.5 Demonstration and Testing", 2)
    add_paragraph(
        doc,
        "Demonstration. The implemented interface currently includes a homepage with search and categories, a post detail page with visible comments, an author dashboard for article creation and editing, and an administrator dashboard for moderation tasks."
    )
    add_paragraph(
        doc,
        "Automated verification was completed on the backend using PHPUnit. The current feature test file covers registration, author approval, publication filtering, authorization boundaries, image upload persistence, admin preview access, and comment creation. In addition, a production frontend build was executed successfully to confirm that the React application compiles without errors."
    )
    add_bullets(
        doc,
        [
            "Backend test result: 9 tests passed with 44 assertions.",
            "Frontend build result: production build compiled successfully.",
            "Manual validation: author image upload, image rendering, and moderation flow were checked during final debugging.",
        ],
    )

    add_heading(doc, "1.6 Challenges and Solutions", 2)
    add_paragraph(
        doc,
        "A key challenge during the final integration stage was image handling. Upload requests initially failed because the shared Axios client forced JSON content headers even for multipart form submissions. This was corrected by allowing FormData requests to set their own multipart boundary automatically."
    )
    add_paragraph(
        doc,
        "A second issue appeared after upload success: image files existed on disk but were not rendered in the browser. The root cause was a mismatch between generated image URLs, storage exposure, and route behavior. The solution was to normalize frontend image paths and introduce a dedicated backend media-serving route that reliably returns uploaded files."
    )
    add_paragraph(
        doc,
        "Beyond media handling, the project also required careful role-based authorization. Laravel middleware and guarded endpoints were used to ensure that normal users cannot create posts, pending authors cannot publish without approval, and administrators alone can approve authors or moderate pending content."
    )

    add_heading(doc, "1.7 Conclusion", 2)
    add_paragraph(
        doc,
        "MeBlog satisfies the core aims of a modern web technologies assignment by combining a client-side SPA with a framework-based backend, relational persistence, authentication, moderation logic, and automated testing. The resulting system is small enough to understand clearly but rich enough to demonstrate practical concerns such as access control, asynchronous API communication, file upload handling, and workflow-based publishing."
    )
    add_paragraph(
        doc,
        "Future enhancements could include richer text editing, image previews before upload, pagination controls for comments, stronger frontend test coverage, deployment documentation, and analytics for authors and administrators."
    )

    doc.add_page_break()

    add_heading(doc, "Task 2 - Testing and Discussion", 1)
    add_heading(doc, "2.1 Testing", 2)
    add_paragraph(
        doc,
        "The current project includes a backend feature test suite implemented in PHPUnit. These tests exercise the API directly and validate both expected success paths and access-control restrictions. Because the repository does not currently include React Testing Library or Jest test files, the frontend validation section below reflects build verification and manual browser testing rather than automated component tests."
    )
    add_paragraph(doc, "Implemented feature tests.")
    add_bullets(
        doc,
        [
            "User registration: verifies that a standard user can register successfully and receive an authentication token.",
            "Author registration workflow: verifies that author sign-up starts in an author_pending state.",
            "Admin author approval: verifies that an administrator can approve a pending author account.",
            "Published post listing: verifies that public post queries return only published content.",
            "Unauthorized post creation: verifies that a normal user is forbidden from author-only content creation.",
            "Author draft creation: verifies that an approved author can create a draft post successfully.",
            "Image upload persistence: verifies that the uploaded file is stored and its path is persisted in the database.",
            "Admin preview of pending post: verifies that an administrator can inspect pending content before review.",
            "Authenticated comment creation: verifies that a signed-in user can create a comment on a published post.",
            "Registration tests cover both standard users and author candidates, ensuring that the moderation workflow starts at account creation time.",
            "Authorization tests confirm that only approved authors can create content and that only administrators can approve authors or moderate pending posts.",
            "The image persistence test verifies an important full-stack behavior: uploaded media is saved on disk and its path is persisted in the database.",
            "Comment creation testing confirms that authenticated interaction is available on published content.",
        ],
    )
    add_paragraph(
        doc,
        "Additional manual checks were completed on the running application. These included registering as an author candidate, approving the account from the admin dashboard, creating and submitting a post, approving it as an administrator, and confirming that the uploaded cover image renders correctly on the public homepage and detail page."
    )
    add_paragraph(
        doc,
        "The frontend was also validated with npm run build, which completed successfully and confirmed that the production bundle can be generated without syntax or import errors."
    )

    add_heading(doc, "2.2 Discussion: Comparing Single Page Applications with Traditional Web Development Methodologies", 2)
    add_paragraph(
        doc,
        "Single Page Applications (SPAs) differ from traditional multi-page web applications because the browser loads one shell page and then updates content dynamically through JavaScript and API calls. In a traditional PHP page-per-request model, each navigation or form action usually returns a new HTML response from the server. SPAs move much of this interaction to the client and rely on APIs to exchange data in the background."
    )
    add_paragraph(
        doc,
        "In MeBlog, React was used on the frontend to create a smoother user experience for searching, navigating between posts, switching dashboards, and updating content after login or moderation events. The interface can update only the parts that change, which generally feels faster and more interactive than a full page reload approach."
    )
    add_paragraph(
        doc,
        "However, SPAs also introduce additional complexity. The developer must manage client-side state, route protection, asynchronous data loading, and error handling explicitly. This was visible in MeBlog when debugging image uploads and media rendering. In a more traditional application, the form post and HTML response would have been easier to trace in one place; in the SPA architecture, the issue had to be tracked through React form state, Axios headers, API validation, and backend file serving."
    )
    add_paragraph(
        doc,
        "Framework support is one of the reasons modern SPAs remain attractive despite this complexity. React encourages the interface to be split into reusable components, while Laravel provides a clean server-side structure for controllers, request validation, middleware, and database models. The combination used in this assignment reflects a common professional pattern: JavaScript-driven presentation on the client and a RESTful API on the server."
    )
    add_paragraph(
        doc,
        "State management is also more visible in SPAs. In MeBlog, search terms, pagination, authenticated user state, comment lists, and role-gated navigation all depend on client-side updates. React hooks make this manageable, but the responsibility still sits on the frontend rather than on server-rendered templates. This changes how developers think about user flows and how they test them."
    )
    add_paragraph(
        doc,
        "CRUD operations also feel different in SPA development. Creating a post, approving a pending article, or adding a comment is done through asynchronous HTTP requests rather than a complete navigation cycle. The benefit is a more responsive interface, but the cost is greater reliance on clean API contracts and independent endpoint testing."
    )
    add_paragraph(
        doc,
        "This is why API testing is particularly important. In a decoupled architecture, failures can come from the client, the server, or the interface between them. PHPUnit tests in this project helped confirm that the backend behaved correctly even before the browser-facing debugging was complete. Once the image endpoint and request headers were fixed, the API test coverage made it easier to trust the server-side rules and focus only on the remaining integration gap."
    )
    add_paragraph(
        doc,
        "SPAs do have trade-offs. Search engine optimization can be more difficult when content is heavily client-rendered, and the initial JavaScript bundle can be larger than a basic server-rendered page. For a project like MeBlog, these drawbacks are acceptable because the assignment focus is on interactivity, modularity, and API integration rather than SEO-first publishing."
    )
    add_paragraph(
        doc,
        "Overall, this assignment shows that SPA development offers a better user experience and a more modular frontend architecture, but it also demands stronger discipline in testing, version control, API design, and debugging. Traditional web development remains simpler for smaller content sites, while SPA-based development becomes more compelling when richer interactions and role-based workflows are required."
    )

    add_heading(doc, "References", 1)
    refs = [
        "[1] A. Banks and E. Porcello, Learning React, 2nd ed., O'Reilly Media, 2020.",
        "[2] Laravel Documentation, 'Laravel 12.x Documentation,' Laravel, accessed 2026.",
        "[3] React Documentation, 'React Developer Documentation,' Meta, accessed 2026.",
        "[4] MDN Web Docs, 'Using Fetch and HTTP Requests in Web Applications,' Mozilla, accessed 2026.",
        "[5] K. C. Dodds, 'Write tests. Not too many. Mostly integration,' kentcdodds.com, 2019.",
        "[6] M. Hartl, Ruby on Rails Tutorial, 6th ed., Addison-Wesley, 2020.",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
